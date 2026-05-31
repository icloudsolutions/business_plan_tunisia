"""
Professional Word feasibility report (VIPA-aligned).

generate_word_report(plan_data, output_path)

plan_data keys:
  inputs, results, yearly (optional), title, plan_id,
  logo_path, sector, promoter, cabinet, market_study, swot
"""

from __future__ import annotations

import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from bp_calc.capex import annual_depreciation_schedule
from bp_calc.engine import HORIZON
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from tasks.chart_generator import generate_all_charts
from tasks.export_excel import (
    YEAR_LABELS,
    ReportContext,
    _categorize_investments,
    _loan_annual_rows,
    _parse_plan_data,
    _product_names,
    _safe_div,
    build_plan_data,
)

CLR_HEADER = "003366"
CLR_H2 = "336699"
CLR_SUBTOTAL = "D3D3D3"
CLR_SWOT = {
    "forces": "C8E6C9",
    "weaknesses": "FFE0B2",
    "opportunities": "BBDEFB",
    "threats": "FFCDD2",
}

# Static table of contents (always rendered; Word TOC fields are often empty until updated).
WORD_REPORT_TOC: list[tuple[str, int]] = [
    ("SECTION 1 — PRESENTATION DU PROJET", 1),
    ("SECTION 2 — ETUDE DE MARCHE", 1),
    ("2.1 Situation actuelle du secteur en Tunisie", 2),
    ("2.2 Clientele cible", 2),
    ("2.3 Concurrence", 2),
    ("2.4 Strategie commerciale", 2),
    ("SECTION 3 — ANALYSE SWOT", 1),
    ("SECTION 4 — INVESTISSEMENT ET FINANCEMENT", 1),
    ("SECTION 5 — ANALYSE DE RENTABILITE (tableaux detailles)", 1),
    ("SECTION 6 — ETAT DE RESULTAT PREVISIONNEL", 1),
    ("SECTION 7 — GRAPHIQUES", 1),
    ("SECTION 8 — INDICATEURS DE RENTABILITE", 1),
    ("Conclusion", 2),
]


def _fmt(n: float | None, *, digits: int = 0) -> str:
    if n is None:
        return "—"
    return f"{n:,.{digits}f}".replace(",", " ")


def _pct(n: float | None) -> str:
    if n is None:
        return "—"
    return f"{n * 100:.2f} %"


def _shade_cell(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _style_header_row(table, ncols: int) -> None:
    for i in range(ncols):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, CLR_HEADER)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.name = "Arial"
                r.font.size = Pt(10)


def _style_total_row(table, row_idx: int, ncols: int) -> None:
    for i in range(ncols):
        cell = table.rows[row_idx].cells[i]
        _shade_cell(cell, CLR_SUBTOTAL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def _add_red_placeholder(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"[{text}]")
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Arial"
        if level == 1:
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(51, 102, 153)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(11)


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    *,
    caption: str | None = None,
    note: str | None = None,
    total_row: bool = False,
) -> None:
    if caption:
        cap = doc.add_paragraph(caption)
        for r in cap.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(11)
    if not rows:
        rows = [["—", "—"]]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    _style_header_row(table, len(headers))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < len(headers):
                table.rows[ri + 1].cells[ci].text = str(val)
    if total_row and len(rows) > 0:
        _style_total_row(table, len(rows), len(headers))
    if note:
        np = doc.add_paragraph(note)
        np.paragraph_format.space_before = Pt(4)
        for r in np.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(102, 102, 102)
    doc.add_paragraph()


def _add_year_table(
    doc: Document,
    title: str,
    row_labels: list[str],
    data_rows: list[list[float | str]],
    *,
    caption: str | None = None,
    pct_rows: set[int] | None = None,
) -> None:
    headers = ["Poste", *YEAR_LABELS]
    rows = []
    for label, vals in zip(row_labels, data_rows):
        row = [label]
        for i, v in enumerate(vals[:HORIZON]):
            if pct_rows and label in pct_rows:
                row.append(_pct(float(v)) if isinstance(v, (int, float)) else str(v))
            else:
                row.append(_fmt(float(v)) if isinstance(v, (int, float)) else str(v))
        rows.append(row)
    _add_table(doc, headers, rows, caption=caption or title)


def _enable_update_fields_on_open(doc: Document) -> None:
    """Ask Word to refresh field codes (TOC, PAGE) when the document opens."""
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _add_toc_line(doc: Document, title: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.6 * (level - 1))
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(12 if level == 1 else 10)
    run.bold = level == 1
    if level == 1:
        run.font.color.rgb = RGBColor(0, 51, 102)
    p.add_run("\t")


def _insert_toc(doc: Document) -> None:
    """Sommaire page after cover: static entries (always visible) + Word TOC field."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("SOMMAIRE")
    tr.bold = True
    tr.font.name = "Arial"
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()

    for entry_title, level in WORD_REPORT_TOC:
        _add_toc_line(doc, entry_title, level)

    hint = doc.add_paragraph(
        "Numeros de page : sous Microsoft Word, inserer une table des matieres automatique "
        "(References > Table des matieres) ou mettre a jour les champs si vous en ajoutez une."
    )
    for r in hint.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor(102, 102, 102)
    doc.add_page_break()


def _setup_header_footer(doc: Document, project_name: str) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        hp.add_run(project_name).bold = True
        hp.add_run("\t\t\t")
        run = hp.add_run("STRICTEMENT CONFIDENTIEL")
        run.bold = True
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run("Page ")
        _add_field(fp, "PAGE")
        fp.add_run(" sur ")
        _add_field(fp, "NUMPAGES")


def _add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def _add_centered_picture(doc: Document, path: str, width_in: float = 6.0) -> None:
    if not Path(path).is_file():
        return
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_figure_block(
    doc: Document,
    path: str | None,
    *,
    figure_num: int,
    title: str,
    comment: str | None = None,
    width_in: float = 6.0,
) -> None:
    cap = doc.add_paragraph(f"Figure {figure_num} — {title}")
    for r in cap.runs:
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(11)
    if path and Path(path).is_file():
        _add_centered_picture(doc, path, width_in)
    else:
        _add_red_placeholder(doc, f"[Graphique {figure_num} — donnees indisponibles]")
    lines = ["Source : moteur de calcul bp_calc (donnees du plan valide)."]
    if comment:
        lines.append(comment)
    note = doc.add_paragraph(" ".join(lines))
    for r in note.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(102, 102, 102)
    doc.add_paragraph()


class WordReportBuilder:
    def __init__(self, plan_data: dict):
        self.plan_data = plan_data
        self.ctx = _parse_plan_data(plan_data)
        self.market = plan_data.get("market_study") or {}
        self.swot = plan_data.get("swot") or {}
        self.sector = plan_data.get("sector") or plan_data.get("activity") or "—"
        self.promoter = plan_data.get("promoter") or self.ctx.company
        self.cabinet = plan_data.get("cabinet") or ""
        self.logo_path = plan_data.get("logo_path")

    def _has_market(self) -> bool:
        return bool(self.market) and any(
            self.market.get(k)
            for k in (
                "sector_overview",
                "ipc_table",
                "trends",
                "market_size",
                "clientele",
                "competitors",
                "commercial_strategy",
            )
        )

    def build_cover(self, doc: Document) -> None:
        if self.logo_path and Path(self.logo_path).is_file():
            try:
                doc.add_picture(self.logo_path, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                _add_red_placeholder(doc, "LOGO — chemin invalide ou format non supporte")
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("[Emplacement logo societe]")
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(f"Business Plan — {self.ctx.company}")
        r.bold = True
        r.font.size = Pt(22)
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor(0, 51, 102)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = sub.add_run(f"Etude de faisabilite — {self.sector}")
        rs.font.size = Pt(14)
        rs.font.name = "Arial"

        doc.add_paragraph()
        for line in [
            f"Promoteur : {self.promoter}",
            f"Date : {datetime.now().strftime('%d/%m/%Y')}",
            f"Cabinet conseil : {self.cabinet or '—'}",
            f"Reference : {self.ctx.plan_id[:8].upper() if self.ctx.plan_id else '—'}",
        ]:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(11)

        conf = doc.add_paragraph()
        conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = conf.add_run("STRICTEMENT CONFIDENTIEL")
        rc.bold = True
        rc.font.color.rgb = RGBColor(180, 0, 0)
        doc.add_page_break()

    def section1_presentation(self, doc: Document) -> None:
        _add_heading(doc, "SECTION 1 — PRESENTATION DU PROJET", 1)
        fin = self.ctx.inputs.financing
        inv = self.ctx.results.totalInvestment
        _add_table(
            doc,
            ["Rubrique", "Valeur"],
            [
                ["Type du projet", str(self.plan_data.get("project_type") or "Creation")],
                ["Activite", str(self.sector)],
                ["Cadre juridique", self.ctx.inputs.company.legalForm],
                ["Implantation", str(self.plan_data.get("site") or "[A completer — implantation]")],
                ["Capital social (DT)", _fmt(inv * fin.equityRatio)],
                [
                    "Participation",
                    str(self.plan_data.get("shareholding") or "100 % tunisienne"),
                ],
                ["Cout du projet (DT)", _fmt(inv)],
                [
                    "Nombre d'emplois",
                    str(
                        sum(p.headcount for p in self.ctx.inputs.plAssumptions.personnel)
                        or "[A completer]"
                    ),
                ],
                [
                    "Produits finis",
                    str(self.plan_data.get("finished_products") or ", ".join(_product_names(self.ctx))),
                ],
            ],
            caption="Fiche projet",
        )
        planning = self.plan_data.get("planning") or [
            ["Etude et montage dossier", "M1-M3", "3 mois"],
            ["Recherche financement", "M2-M6", "5 mois"],
            ["Commande equipements", "M4-M8", "5 mois"],
            ["Demarrage production", "M9-M12", "4 mois"],
        ]
        if isinstance(planning, list) and planning and isinstance(planning[0], dict):
            planning = [
                [p.get("phase", ""), p.get("period", ""), p.get("duration", "")]
                for p in planning
            ]
        _add_table(
            doc,
            ["Phase", "Periode", "Duree"],
            planning if isinstance(planning[0], list) else [["—", "—", "—"]],
            caption="Planning previsionnel de realisation",
        )
        _add_table(
            doc,
            ["Nature de l'avantage", "Description"],
            [
                ["Prime d'investissement", str(self.plan_data.get("prime_invest") or "[A completer]")],
                ["Prime etude / assistance", str(self.plan_data.get("prime_etude") or "[A completer]")],
                ["Prime immateriels", str(self.plan_data.get("prime_immat") or "[A completer]")],
            ],
            caption="Avantages financiers (Code des incitations)",
        )

    def section2_market(self, doc: Document) -> None:
        _add_heading(doc, "SECTION 2 — ETUDE DE MARCHE", 1)
        if not self._has_market():
            _add_red_placeholder(
                doc,
                "ETUDE DE MARCHE INCOMPLETE — Renseigner plan_data['market_study'] "
                "(secteur, IPC, taille marche, clientele, concurrence, strategie commerciale)"
            )

        _add_heading(doc, "2.1 Situation actuelle du secteur en Tunisie", 2)
        ipc = self.market.get("ipc_table")
        if ipc and isinstance(ipc, dict):
            years = ipc.get("years") or []
            categories = ipc.get("categories") or {}
            headers = ["Categorie", *[str(y) for y in years]]
            rows = [[cat, *[str(v) for v in vals]] for cat, vals in categories.items()]
            _add_table(doc, headers, rows, caption="Indices des prix a la consommation (IPC)")
        else:
            _add_red_placeholder(doc, "Tableau IPC par categorie et annees — a completer")
        trends = self.market.get("trends") or self.market.get("sector_overview")
        if trends:
            _add_body(doc, str(trends))
        else:
            _add_red_placeholder(doc, "Tendances du marche et comportement consommateurs — a rediger")
        msize = self.market.get("market_size")
        if msize:
            if isinstance(msize, dict):
                _add_table(
                    doc,
                    ["Indicateur", "Valeur"],
                    [[k, str(v)] for k, v in msize.items()],
                    caption="Taille du marche cible",
                )
            else:
                _add_body(doc, str(msize))
        else:
            _add_red_placeholder(doc, "Taille du marche et part estimee du projet — a completer")

        _add_heading(doc, "2.2 Clientele cible", 2)
        cli = self.market.get("clientele")
        if cli:
            _add_body(doc, str(cli) if isinstance(cli, str) else "")
            if isinstance(cli, dict):
                _add_table(doc, ["Segment", "Detail"], [[k, str(v)] for k, v in cli.items()])
        else:
            _add_red_placeholder(
                doc, "Segmentation B2C/B2B, GMS, export — zones geographiques — a completer"
            )

        _add_heading(doc, "2.3 Concurrence", 2)
        comp = self.market.get("competitors")
        if comp and isinstance(comp, list):
            _add_table(
                doc,
                ["Concurrent", "Part marche", "Forces", "Faiblesses"],
                [
                    [
                        c.get("name", "—"),
                        c.get("share", "—"),
                        c.get("strengths", "—"),
                        c.get("weaknesses", "—"),
                    ]
                    for c in comp
                    if isinstance(c, dict)
                ],
                caption="Analyse concurrentielle",
            )
        else:
            _add_red_placeholder(doc, "Tableau concurrents directs/indirects — a completer")

        _add_heading(doc, "2.4 Strategie commerciale", 2)
        strat = self.market.get("commercial_strategy")
        if strat:
            if isinstance(strat, dict):
                _add_table(doc, ["Element", "Detail"], [[k, str(v)] for k, v in strat.items()])
            else:
                _add_body(doc, str(strat))
        else:
            disc = self.ctx.inputs.plAssumptions.commercialDiscount
            ops = self.ctx.inputs.operations
            _add_body(
                doc,
                f"Prix de vente unitaire retenu : {_fmt(ops.salePrice)} DT. "
                f"Ristourne commerciale : {_pct(disc)}. "
                "[A completer — canaux de distribution et gammes prix]"
            )

    def section3_swot(self, doc: Document) -> None:
        _add_heading(doc, "SECTION 3 — ANALYSE SWOT", 1)
        defaults = {
            "forces": [
                "Innovation au niveau de la valorisation des produits",
                "Variety des produits",
                "Rentabilite financiere satisfaisante",
                "Produits destines a toutes categories de consommateurs",
            ],
            "weaknesses": ["Concurrence avec les produits importes"],
            "opportunities": [
                "Absence d'une forte concurrence locale sur certains segments",
                "Sensibilite croissante des consommateurs a la qualite",
            ],
            "threats": ["Entree de nouveaux competiteurs (cout du projet modere)"],
        }
        swot_data = {
            "forces": self.swot.get("forces") or defaults["forces"],
            "weaknesses": self.swot.get("weaknesses") or defaults["weaknesses"],
            "opportunities": self.swot.get("opportunities") or defaults["opportunities"],
            "threats": self.swot.get("threats") or defaults["threats"],
        }
        if not self.swot:
            _add_red_placeholder(doc, "SWOT generee depuis modele VIPA — valider ou completer plan_data['swot']")

        def _cell_text(items) -> str:
            if isinstance(items, str):
                return items
            return "\n".join(f"• {x}" for x in items)

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        cells = [
            ("FORCES", "forces", swot_data["forces"]),
            ("FAIBLESSES", "weaknesses", swot_data["weaknesses"]),
            ("OPPORTUNITES", "opportunities", swot_data["opportunities"]),
            ("MENACES", "threats", swot_data["threats"]),
        ]
        positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
        for (title, key, content), (ri, ci) in zip(cells, positions):
            cell = table.rows[ri].cells[ci]
            _shade_cell(cell, CLR_SWOT[key])
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(f"{title}\n")
            r.bold = True
            p.add_run(_cell_text(content))
        doc.add_paragraph()

    def section4_investment(self, doc: Document) -> None:
        _add_heading(doc, "SECTION 4 — INVESTISSEMENT ET FINANCEMENT", 1)
        buckets = _categorize_investments(self.ctx)
        total = self.ctx.results.totalInvestment
        for cat_label, key in [
            ("Immobilisations incorporelles", "incorporel"),
            ("Agencements et amenagement", "agencement"),
            ("Materiels industriels", "industriel"),
            ("Materiel de transport", "transport"),
            ("Mobilier de bureau", "bureau"),
            ("Frais preliminaires", "preliminaires"),
        ]:
            items = buckets.get(key, [])
            rows = []
            for name, amt, _life, _rate in items:
                rows.append([name, _fmt(amt), _pct(_safe_div(amt, total))])
            if not rows:
                rows = [["—", "0", "0 %"]]
            _add_table(
                doc,
                ["Designation", "Valeur DT", "% total"],
                rows,
                caption=cat_label,
            )
        bfr0 = self.ctx.results.bfr.years[0] if self.ctx.results.bfr.years else 0
        wc = self.ctx.inputs.workingCapital
        _add_body(
            doc,
            f"BFR initial (An 1) : {_fmt(bfr0)} DT — delais clients {wc.clientPaymentDays} j, "
            f"stock PF {wc.finishedGoodsStockDays} j, stock MP {wc.rawMaterialStockMonths:.1f} mois."
        )
        _add_table(
            doc,
            ["Poste", "Valeur DT", "%"],
            [
                ["Total investissement", _fmt(total), "100 %"],
                ["BFR initial", _fmt(bfr0), _pct(_safe_div(bfr0, total + bfr0))],
                ["Investissement + BFR", _fmt(total + bfr0), "—"],
            ],
            caption="Tableau 4 — Total investissement et BFR",
            note="BFR = besoin en fonds de roulement initial (An 1).",
            total_row=True,
        )
        fin = self.ctx.inputs.financing
        eq = total * fin.equityRatio
        debt = fin.loan.amount or total * fin.debtRatio
        _add_table(
            doc,
            ["Source", "Valeur DT", "%"],
            [
                ["Fonds propres", _fmt(eq), _pct(fin.equityRatio)],
                ["Credit moyen terme", _fmt(debt), _pct(fin.debtRatio)],
                ["Subventions", _fmt(float(self.plan_data.get("subsidies") or 0)), "—"],
            ],
            caption="Tableau 5 — Schema de financement",
            note="Repartition fonds propres, credit moyen terme et subventions eventuelles.",
        )
        loan = fin.loan
        _add_table(
            doc,
            ["Parametre", "Valeur"],
            [
                ["Montant emprunt (DT)", _fmt(debt)],
                ["Duree (ans)", str(loan.years)],
                ["Taux d'interet", _pct(loan.rate)],
                ["Differe principal (mois)", str(loan.graceMonthsPrincipal)],
                ["TMM + spread", str(self.plan_data.get("loan_spread") or "TMM + 4 %")],
            ],
            caption="Conditions du credit",
        )

    def section5_profitability_tables(self, doc: Document) -> None:
        _add_heading(doc, "SECTION 5 — ANALYSE DE RENTABILITE (tableaux detailles)", 1)
        products = _product_names(self.ctx)
        rev = self.ctx.rev()
        n = max(1, len(products))
        per_prod = [[rev[y] / n for y in range(HORIZON)] for _ in products]
        _add_year_table(
            doc,
            "Chiffre d'affaires par produit (DT)",
            products,
            per_prod,
        )
        growth = [0.0] + [_safe_div(rev[y] - rev[y - 1], rev[y - 1]) for y in range(1, HORIZON)]
        _add_year_table(doc, "Taux croissance CA", ["Croissance"], [growth])
        _add_year_table(doc, "CA net total", ["Total"], [rev])

        mp = self.ctx.purchase_mp()
        _add_year_table(doc, "Achats matieres premieres", ["Total MP"], [mp])
        packaging, other = [], []
        for y in range(HORIZON):
            yr = self.ctx.yearly[y] if y < len(self.ctx.yearly) else {}
            cogs = float(yr.get("cogs", 0))
            pack = self.ctx.inputs.operations.packagingCost * (
                self.ctx.results.qtySold.years[y] if y < len(self.ctx.results.qtySold.years) else 0
            )
            packaging.append(pack)
            other.append(max(0.0, cogs - mp[y] - pack))
        _add_year_table(
            doc,
            "Autres approvisionnements",
            ["Emballages / aromes / energie"],
            [packaging],
        )
        _add_year_table(doc, "Autres achats", ["Autres"], [other])
        total_achats = [[mp[y] + packaging[y] + other[y] for y in range(HORIZON)]]
        _add_year_table(doc, "Total achats consommes", ["Total"], total_achats)

        margin = []
        for y in range(HORIZON):
            yr = self.ctx.yearly[y] if y < len(self.ctx.yearly) else {}
            margin.append(float(yr.get("grossMarginPct", _safe_div(rev[y] - float(yr.get("cogs", 0)), rev[y]))))
        _add_year_table(doc, "Taux de marge brute", ["Marge %"], [margin])

        pers_rows = []
        labels = []
        for p in self.ctx.inputs.plAssumptions.personnel:
            if p.role.strip() or p.headcount:
                labels.append(p.role)
                pers_rows.append([float(p.headcount)] * HORIZON)
        if labels:
            _add_year_table(doc, "Effectifs par fonction", labels, pers_rows)
        total_sal = sum(p.headcount * p.annualSalary for p in self.ctx.inputs.plAssumptions.personnel)
        cnss = total_sal * 0.19
        _add_year_table(
            doc,
            "Charges de personnel",
            ["Salaire brut", "CNSS", "Total"],
            [[total_sal] * HORIZON, [cnss] * HORIZON, [total_sal + cnss] * HORIZON],
        )

        dep_labels, dep_rows = [], []
        for eq in self.ctx.inputs.investments.equipment:
            if eq.cost <= 0:
                continue
            life = max(1, eq.usefulLifeYears)
            annual = eq.cost / life
            line = [0.0] * HORIZON
            start = max(0, min(HORIZON - 1, eq.acquisitionYear - 1))
            for off in range(life):
                yi = start + off
                if yi < HORIZON:
                    line[yi] = annual
            dep_labels.append(f"{eq.name[:35]} ({_pct(1/life)})")
            dep_rows.append(line)
        dep_sched = annual_depreciation_schedule(self.ctx.inputs)
        dep_labels.append("Total DAP")
        dep_rows.append(dep_sched)
        _add_year_table(doc, "Dotations aux amortissements", dep_labels, dep_rows)

        _add_year_table(
            doc,
            "Autres charges d'exploitation",
            ["Marketing", "Distribution", "Autres"],
            [self.ctx.mkt(), self.ctx.dist(), [float(self.ctx.yearly[y].get("otherOpex", 0)) for y in range(HORIZON)]],
        )

        loan_rows = _loan_annual_rows(self.ctx)
        _add_year_table(
            doc,
            "Tableau amortissement emprunt",
            [
                "Capital restant du (debut)",
                "Interets",
                "Remboursement principal",
                "Service de la dette",
                "Capital restant du (fin)",
            ],
            [
                [r["opening"] for r in loan_rows],
                [r["interest"] for r in loan_rows],
                [r["principal"] for r in loan_rows],
                [r["service"] for r in loan_rows],
                [r["closing"] for r in loan_rows],
            ],
        )

    def section6_pl(self, doc: Document) -> None:
        _add_heading(doc, "SECTION 6 — ETAT DE RESULTAT PREVISIONNEL", 1)
        rev_total = sum(self.ctx.rev()) or 1
        pl_rows = [
            ("Chiffre d'affaires HT", "revenue"),
            ("Achats consommes", "cogs"),
            ("Charges de personnel", "personnel"),
            ("Frais distribution", "distribution"),
            ("Frais marketing", "marketing"),
            ("TVA nette", "vat"),
            ("EBE", "ebe"),
            ("Dotations amortissement", "depreciation"),
            ("Resultat exploitation", "_ebit"),
            ("Charges financieres", "interest"),
            ("Resultat avant IS", "_pretax"),
            ("Impot sur societes", "tax"),
            ("Resultat net", "netProfit"),
        ]
        headers = ["Poste", *YEAR_LABELS, "% Revenu"]
        rows = []
        for label, key in pl_rows:
            vals = []
            total_line = 0.0
            for y in range(HORIZON):
                yr = self.ctx.yearly[y] if y < len(self.ctx.yearly) else {}
                if key == "_ebit":
                    v = float(yr.get("ebe", 0)) - float(yr.get("depreciation", 0))
                elif key == "_pretax":
                    v = (
                        float(yr.get("ebe", 0))
                        - float(yr.get("depreciation", 0))
                        - float(yr.get("interest", 0))
                    )
                elif key == "netProfit":
                    v = self.ctx.net()[y]
                else:
                    v = float(yr.get(key, 0))
                vals.append(v)
                total_line += v
            row = [label] + [_fmt(v) for v in vals] + [_pct(_safe_div(total_line, rev_total))]
            rows.append(row)
        _add_table(
            doc,
            headers,
            rows,
            caption="Tableau 6 — Compte de resultat previsionnel (DT)",
            note="Montants en dinars tunisiens ; colonne % Revenu = part du poste sur le total CA sur 7 ans.",
        )

    def section7_charts(self, doc: Document, chart_dir: Path) -> None:
        _add_heading(doc, "SECTION 7 — GRAPHIQUES", 1)
        ebit = []
        for y in range(HORIZON):
            yr = self.ctx.yearly[y] if y < len(self.ctx.yearly) else {}
            ebit.append(float(yr.get("ebe", 0)) - float(yr.get("depreciation", 0)))
        products = _product_names(self.ctx)
        rev = self.ctx.rev()
        n = max(1, len(products))
        by_product = {p: [rev[y] / n for y in range(HORIZON)] for p in products}

        chart_data = {
            "revenue": self.ctx.rev(),
            "net_profit": self.ctx.net(),
            "ebit": ebit,
            "cumulative_treasury": self.ctx.cum_treasury(),
            "total_investment": self.ctx.results.totalInvestment,
            "drci_years": self.ctx.ind.drciYears,
            "labels": ["Achats matieres", "Personnel", "DAP", "Autres charges"],
            "values": [],
            "by_product": by_product,
        }
        personnel = sum(p.headcount * p.annualSalary for p in self.ctx.inputs.plAssumptions.personnel) * HORIZON
        achats = sum(self.ctx.purchase_mp())
        dap = sum(self.ctx.dep())
        autres = sum(
            float(self.ctx.yearly[y].get("otherOpex", 0))
            + self.ctx.mkt()[y]
            + self.ctx.dist()[y]
            for y in range(min(HORIZON, len(self.ctx.yearly)))
        )
        chart_data["values"] = [achats, personnel, dap, autres]

        paths = generate_all_charts(chart_data, chart_dir)
        figures = [
            (
                "results_evolution",
                "Evolution du chiffre d'affaires, du resultat d'exploitation et du resultat net",
                "Courbes sur 7 ans ; resultat d'exploitation = EBE − dotations.",
            ),
            (
                "cumulative_treasury",
                "Tresorerie cumulee et delai de recuperation (DRCI)",
                "Barres vertes/rouges selon signe ; trait pointille = investissement initial.",
            ),
            (
                "cost_structure",
                "Structure des couts (moyenne sur 7 ans)",
                "Repartition achats, personnel, amortissements et autres charges.",
            ),
            (
                "ca_by_product",
                "Chiffre d'affaires par produit",
                "Histogramme empile par annee et par ligne de produit.",
            ),
        ]
        for i, (key, title, comment) in enumerate(figures, start=1):
            _add_figure_block(
                doc,
                paths.get(key),
                figure_num=i,
                title=title,
                comment=comment,
            )

    def section8_indicators(self, doc: Document) -> None:
        _add_heading(doc, "SECTION 8 — INDICATEURS DE RENTABILITE", 1)
        ind = self.ctx.ind
        _add_table(
            doc,
            ["Indicateur", "Valeur", "Interpretation"],
            [
                [
                    "VAN (DT)",
                    _fmt(ind.van),
                    "Positive — projet economiquement viable"
                    if ind.van >= 0
                    else "Negative — ajuster hypotheses",
                ],
                [
                    "TRI",
                    _pct(ind.tri),
                    "Superieur au cout du capital" if ind.tri and ind.tri > ind.discountRate else "A analyser",
                ],
                [
                    "DRCI (ans)",
                    _fmt(ind.drciYears, digits=1) if ind.drciYears else "—",
                    "Delai de recuperation de l'investissement",
                ],
                ["Taux d'actualisation", _pct(ind.discountRate), "Hypothese retenue"],
            ],
            caption="Tableau 7 — Synthese des indicateurs de rentabilite",
            note="VAN et TRI calcules avec le taux d'actualisation indique ; DRCI = annee de recuperation du flux cumule.",
        )
        disc = ind.discountRate
        ocf = self.ctx.ocf()
        _add_year_table(doc, "Cash-flows exploitation", ["CF exploitation"], [ocf])
        discounted = [v / ((1 + disc) ** (y + 1)) for y, v in enumerate(ocf)]
        _add_year_table(doc, "Cash-flows actualises", ["CF actualise"], [discounted])

        inv = self.ctx.results.totalInvestment
        flows = [-inv, *ocf]
        _add_body(
            doc,
            f"Flux pour TRI : investissement initial {_fmt(inv)} DT, puis CF exploitation sur 7 ans. "
            f"TRI estime : {_pct(ind.tri)}."
        )
        if ind.van >= 0 and self.ctx.results.cashRunwayBreakYear is None:
            conclusion = (
                f"Sous les hypotheses retenues, le projet « {self.ctx.company} » affiche une VAN de "
                f"{_fmt(ind.van)} DT et un TRI de {_pct(ind.tri)}. La tresorerie reste positive sur "
                "l'horizon septennal. Le dossier peut etre presente aux instances d'investissement."
            )
        else:
            conclusion = (
                f"Le projet « {self.ctx.company} » necessite un ajustement des hypotheses "
                f"(VAN {_fmt(ind.van)} DT). Renforcer le BFR, le plan commercial ou le financement."
            )
        _add_heading(doc, "Conclusion", 2)
        _add_body(doc, conclusion)


def generate_word_report(plan_data: dict, output_path: str) -> None:
    """Build the complete Word feasibility report."""
    doc = Document()
    _enable_update_fields_on_open(doc)
    builder = WordReportBuilder(plan_data)
    _setup_header_footer(doc, builder.ctx.company)
    builder.build_cover(doc)
    _insert_toc(doc)
    builder.section1_presentation(doc)
    builder.section2_market(doc)
    builder.section3_swot(doc)
    builder.section4_investment(doc)
    builder.section5_profitability_tables(doc)
    builder.section6_pl(doc)

    chart_dir = tempfile.mkdtemp(prefix="bp_charts_")
    try:
        builder.section7_charts(doc, Path(chart_dir))
        builder.section8_indicators(doc)
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
    finally:
        shutil.rmtree(chart_dir, ignore_errors=True)


def celery_export_word(plan_data: dict, output_path: str) -> str:
    generate_word_report(plan_data, output_path)
    return str(Path(output_path).resolve())


def build_feasibility_docx_from_plan_data(plan_data: dict, export_dir: Path) -> str:
    """Backward-compatible path for Celery export job."""
    ctx = _parse_plan_data(plan_data)
    slug = "".join(c if c.isalnum() else "_" for c in ctx.company)[:40].strip("_") or "projet"
    path = export_dir / f"etude_faisabilite_{ctx.plan_id}_{slug}.docx"
    generate_word_report(plan_data, str(path))
    return str(path.resolve())
