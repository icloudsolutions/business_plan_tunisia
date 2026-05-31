"""Generate a VIPA-style étude de faisabilité (Word DOCX) from plan inputs and results."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from bp_schema.liasse import PlanInputs, PlanResults
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _fmt(n: float | None, *, digits: int = 0) -> str:
    if n is None:
        return "—"
    return f"{n:,.{digits}f}".replace(",", " ")


def _pct(n: float | None) -> str:
    if n is None:
        return "—"
    return f"{n * 100:.2f} %"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def build_feasibility_docx(
    plan_id: str,
    inputs: PlanInputs,
    results: PlanResults,
    *,
    export_dir: Path,
    plan_title: str | None = None,
) -> str:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    company = inputs.company.name.strip() or "Projet"
    title = plan_title or f"Business Plan — {company}"
    today = datetime.now().strftime("%d/%m/%Y")

    # —— Page de garde ——
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ÉTUDE DE FAISABILITÉ")
    run.bold = True
    run.font.size = Pt(22)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(company)
    r2.bold = True
    r2.font.size = Pt(16)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{title}\n")
    meta.add_run(f"Forme juridique : {inputs.company.legalForm}\n")
    meta.add_run(f"Date du rapport : {today}\n")
    meta.add_run(f"Référence dossier : {plan_id[:8].upper()}")

    doc.add_page_break()

    # —— 1. Synthèse ——
    _add_heading(doc, "1. Synthèse exécutive", 1)
    doc.add_paragraph(
        f"Le présent document constitue l'étude de faisabilité du projet « {company} », "
        "établie à partir de la liasse unique et des projections financières sur sept ans. "
        "Il synthétise les investissements, les hypothèses d'exploitation, le plan de financement "
        "et les indicateurs de rentabilité retenus pour la décision d'investissement."
    )

    ind = results.indicators
    _add_table(
        doc,
        ["Indicateur", "Valeur"],
        [
            ["Investissement total (CAPEX)", f"{_fmt(results.totalInvestment)} TND"],
            ["Valeur Actuelle Nette (VAN)", f"{_fmt(ind.van)} TND"],
            ["Taux de Rentabilité Interne (TRI)", _pct(ind.tri)],
            ["Délai de récupération (DRCI)", f"{_fmt(ind.drciYears, digits=1)} ans" if ind.drciYears else "—"],
            ["Bilan prévisionnel équilibré", "Oui" if results.balanceSheetBalanced else "Non"],
            ["BFR cohérent", "Oui" if results.bfrCoherent else "Non"],
            [
                "Horizon de trésorerie positive",
                "Oui"
                if results.cashRunwayBreakYear is None
                else f"Rupture dès l'an {results.cashRunwayBreakYear}",
            ],
        ],
    )

    # —— 2. Présentation ——
    _add_heading(doc, "2. Présentation du projet", 1)
    doc.add_paragraph(
        f"Le projet est porté par une structure de type {inputs.company.legalForm}. "
        "Les hypothèses détaillées figurent dans les sections suivantes et dans la liasse "
        "unique transmise au guichet d'investissement."
    )

    # —— 3. Investissements ——
    _add_heading(doc, "3. Programme d'investissement", 1)
    inv_rows: list[list[str]] = []
    for eq in inputs.investments.equipment:
        if eq.cost > 0 or eq.name.strip():
            inv_rows.append(
                [
                    eq.name,
                    "Incorporel" if eq.assetType == "intangible" else "Corporel",
                    _fmt(eq.cost),
                    str(eq.usefulLifeYears),
                    f"An {eq.acquisitionYear}",
                ]
            )
    for line in inputs.investments.intangible:
        if line.amount > 0:
            inv_rows.append([line.label, "Incorporel", _fmt(line.amount), str(line.usefulLifeYears), "—"])
    for line in inputs.investments.tangible:
        if line.amount > 0:
            inv_rows.append([line.label, "Corporel", _fmt(line.amount), str(line.usefulLifeYears), "—"])

    if inv_rows:
        _add_table(
            doc,
            ["Désignation", "Nature", "Montant (TND)", "Durée amort.", "Mise en service"],
            inv_rows,
        )
    else:
        doc.add_paragraph("Aucun investissement détaillé n'a été renseigné dans la liasse.")
    doc.add_paragraph(f"Total investissement retenu : {_fmt(results.totalInvestment)} TND.")

    # —— 4. Exploitation ——
    _add_heading(doc, "4. Hypothèses d'exploitation", 1)
    ops = inputs.operations
    _add_table(
        doc,
        ["Paramètre", "Valeur"],
        [
            ["Jours ouvrés / an", _fmt(ops.workingDaysPerYear, digits=0)],
            ["Heures / jour", _fmt(ops.hoursPerDay, digits=1)],
            ["Prix de vente unitaire", f"{_fmt(ops.salePrice)} TND"],
            ["Coût matière unitaire", f"{_fmt(ops.rawMaterialCost)} TND"],
            ["Taux de déchet", _pct(ops.wasteRate.value)],
        ],
    )

    # —— 5. Financement ——
    _add_heading(doc, "5. Plan de financement", 1)
    fin = inputs.financing
    loan = fin.loan
    _add_table(
        doc,
        ["Source", "Part / détail"],
        [
            ["Fonds propres", _pct(fin.equityRatio)],
            ["Dette", _pct(fin.debtRatio)],
            ["Montant emprunt retenu", f"{_fmt(loan.amount or results.totalInvestment * fin.debtRatio)} TND"],
            ["Taux d'intérêt", _pct(loan.rate)],
            ["Durée emprunt", f"{loan.years} ans"],
            ["Différé principal", f"{loan.graceMonthsPrincipal} mois"],
        ],
    )

    wc = inputs.workingCapital
    _add_heading(doc, "5.1 Besoin en fonds de roulement", 2)
    _add_table(
        doc,
        ["Poste", "Valeur"],
        [
            ["Délai clients (jours)", str(wc.clientPaymentDays)],
            ["Délai fournisseurs (jours)", str(wc.supplierPaymentDays)],
            ["Stock produits finis (jours)", str(wc.finishedGoodsStockDays)],
            ["Stock matières (mois)", _fmt(wc.rawMaterialStockMonths, digits=1)],
        ],
    )

    # —— 6. Compte de résultat 7 ans ——
    _add_heading(doc, "6. Compte de résultat et activité (7 ans)", 1)
    pl_rows = []
    for y in range(7):
        rev = results.revenue.years[y] if y < len(results.revenue.years) else 0
        np = results.netProfit.years[y] if y < len(results.netProfit.years) else 0
        ocf = results.operatingCashFlow.years[y] if y < len(results.operatingCashFlow.years) else 0
        pl_rows.append([f"An {y + 1}", _fmt(rev), _fmt(np), _fmt(ocf)])
    _add_table(
        doc,
        ["Exercice", "CA HT (TND)", "Résultat net (TND)", "CF exploitation (TND)"],
        pl_rows,
    )

    # —— 7. Trésorerie ——
    _add_heading(doc, "7. Trésorerie cumulée et BFR", 1)
    treas_rows = []
    for y in range(7):
        tre = results.cumulativeTreasury.years[y] if y < len(results.cumulativeTreasury.years) else 0
        bfr = results.bfr.years[y] if y < len(results.bfr.years) else 0
        treas_rows.append([f"An {y + 1}", _fmt(tre), _fmt(bfr)])
    _add_table(doc, ["Exercice", "Trésorerie cumulée (TND)", "BFR (TND)"], treas_rows)

    # —— 8. Conclusion ——
    _add_heading(doc, "8. Conclusion et recommandation", 1)
    if ind.van >= 0 and results.balanceSheetBalanced and results.cashRunwayBreakYear is None:
        conclusion = (
            f"Sous les hypothèses retenues, le projet « {company} » présente une VAN positive "
            f"({_fmt(ind.van)} TND) et une structure financière équilibrée sur l'horizon de sept ans. "
            "Le dossier peut être présenté pour examen par les instances compétentes, sous réserve "
            "de validation des pièces justificatives complémentaires."
        )
    else:
        conclusion = (
            f"Le projet « {company} » présente des alertes financières (VAN, trésorerie ou bilan). "
            "Il est recommandé d'ajuster les hypothèses d'exploitation, le BFR ou le plan de financement "
            "avant dépôt définitif du dossier."
        )
    doc.add_paragraph(conclusion)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        "Document généré automatiquement par Business Plan Tunisie — "
        "à compléter par les parties narratives (marché, technique, environnement) si requis par le guichet."
    ).italic = True

    slug = "".join(c if c.isalnum() else "_" for c in company)[:40].strip("_") or "projet"
    path = export_dir / f"etude_faisabilite_{plan_id}_{slug}.docx"
    doc.save(str(path))
    return str(path.resolve())
