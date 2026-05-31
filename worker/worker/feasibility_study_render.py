"""Render VIPA study content to DOCX, PDF and XLSX."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from bp_schema.liasse import PlanInputs, PlanResults
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.shapes import Drawing, String
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from worker.feasibility_narrative import NarrativeContext, year_val
from worker.feasibility_vipa_template import (
    VIPA_SOMMAIRE,
    SectionBlock,
    TableBlock,
    VipaStudyBuilder,
)
from worker.pdf_common import pdf_safe, pdf_table

HORIZON = 7
_CHART_W = 16 * cm
_CHART_H = 7 * cm


def _slug(company: str) -> str:
    return "".join(c if c.isalnum() else "_" for c in company)[:40].strip("_") or "projet"


def _make_builder(
    plan_id: str,
    inputs: PlanInputs,
    results: PlanResults,
    *,
    plan_title: str | None = None,
    extra_inputs: dict | None = None,
) -> VipaStudyBuilder:
    ctx = NarrativeContext(
        plan_id=plan_id,
        inputs=inputs,
        results=results,
        plan_title=plan_title,
    )
    return VipaStudyBuilder(ctx, extra_inputs=extra_inputs)


# —— DOCX ——


def _docx_add_table(doc: Document, block: TableBlock) -> None:
    if block.caption:
        doc.add_paragraph(block.caption)
    table = doc.add_table(rows=1 + len(block.rows), cols=len(block.headers))
    table.style = "Table Grid"
    for i, h in enumerate(block.headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(block.rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()


def _docx_add_section(doc: Document, section: SectionBlock) -> None:
    if section.title:
        doc.add_heading(section.title, level=min(section.level, 3))
    for para in section.paragraphs:
        doc.add_paragraph(para)
    for bullet in section.bullets:
        doc.add_paragraph(bullet, style="List Bullet")
    for tbl in section.tables:
        _docx_add_table(doc, tbl)


def build_feasibility_docx(
    plan_id: str,
    inputs: PlanInputs,
    results: PlanResults,
    *,
    export_dir: Path,
    plan_title: str | None = None,
    extra_inputs: dict | None = None,
) -> str:
    builder = _make_builder(
        plan_id, inputs, results, plan_title=plan_title, extra_inputs=extra_inputs
    )
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for i, line in enumerate(builder.cover_block()):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        if i == 0:
            run.bold = True
            run.font.size = Pt(14)
        elif line == "Etude de faisabilité":
            run.bold = True
            run.font.size = Pt(22)
        elif line.startswith("«"):
            run.bold = True
            run.font.size = Pt(16)

    doc.add_page_break()

    doc.add_heading("SOMMAIRE", level=1)
    doc.add_paragraph()
    current = ""
    for entry in VIPA_SOMMAIRE:
        if entry.section:
            current = entry.section
            p = doc.add_paragraph()
            r = p.add_run(entry.section)
            r.bold = True
        if entry.subsection:
            line = f"    {entry.subsection}"
            tab = doc.add_paragraph()
            tab.add_run(f"{line}\t{entry.page}")
    doc.add_paragraph()
    doc.add_page_break()

    for sec in builder.all_sections():
        _docx_add_section(doc, sec)
        if sec.level == 1:
            doc.add_paragraph()

    foot = doc.add_paragraph()
    foot.add_run(
        f"Document généré par Business Plan Tunisie — {datetime.now().strftime('%d/%m/%Y %H:%M')}. "
        "Les mentions « à compléter » doivent être renseignées dans la liasse ou en annexe."
    ).italic = True

    path = export_dir / f"etude_faisabilite_{plan_id}_{_slug(builder.company)}.docx"
    doc.save(str(path))
    return str(path.resolve())


# —— PDF ——


def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle(
            "H1",
            parent=styles["Heading1"],
            fontSize=14,
            spaceBefore=14,
            spaceAfter=8,
            textColor=colors.HexColor("#1E3A5F"),
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=styles["Heading2"],
            fontSize=11,
            spaceBefore=10,
            spaceAfter=6,
            textColor=colors.HexColor("#1E3A5F"),
        ),
        "body": ParagraphStyle("Body", parent=styles["Normal"], fontSize=9, leading=12, spaceAfter=5),
        "cover": ParagraphStyle(
            "Cover",
            parent=styles["Heading1"],
            fontSize=18,
            alignment=1,
            textColor=colors.HexColor("#1E3A5F"),
        ),
        "center": ParagraphStyle("Center", parent=styles["Normal"], alignment=1, fontSize=10),
    }


def _pdf_bar_chart(title: str, labels: list[str], values: list[float], *, fill: str = "#4F46E5") -> Drawing:
    d = Drawing(_CHART_W, _CHART_H)
    d.add(
        String(
            0,
            _CHART_H - 14,
            pdf_safe(title),
            fontSize=10,
            fillColor=colors.HexColor("#1E3A5F"),
        )
    )
    bc = VerticalBarChart()
    bc.x = 30
    bc.y = 20
    bc.height = _CHART_H - 50
    bc.width = _CHART_W - 50
    bc.data = [values]
    bc.categoryAxis.categoryNames = labels
    bc.valueAxis.valueMin = 0
    vmax = max(values) if values and max(values) > 0 else 1
    bc.valueAxis.valueMax = vmax * 1.15
    bc.bars[0].fillColor = colors.HexColor(fill)
    d.add(bc)
    return d


def _pdf_add_table(story: list, block: TableBlock, body_style) -> None:
    if block.caption:
        story.append(Paragraph(pdf_safe(block.caption), body_style))
    data = [block.headers, *block.rows]
    ncols = len(block.headers)
    width = 17 * cm
    col_w = [width / ncols] * ncols
    story.append(pdf_table(data, col_w))
    story.append(Spacer(1, 8))


def build_etude_faisabilite_pdf(
    plan_id: str,
    inputs: PlanInputs,
    results: PlanResults,
    path: Path,
    *,
    plan_title: str | None = None,
    extra_inputs: dict | None = None,
) -> str:
    builder = _make_builder(
        plan_id, inputs, results, plan_title=plan_title, extra_inputs=extra_inputs
    )
    out_path = path / f"etude_faisabilite_{plan_id}_{_slug(builder.company)}.pdf"
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )
    st = _pdf_styles()
    story: list = []

    story.append(Spacer(1, 2.5 * cm))
    for line in builder.cover_block():
        if line:
            story.append(Paragraph(pdf_safe(line), st["cover"] if line == "Etude de faisabilité" else st["center"]))
        else:
            story.append(Spacer(1, 8))
    story.append(PageBreak())

    story.append(Paragraph(pdf_safe("SOMMAIRE"), st["h1"]))
    story.append(Spacer(1, 6))
    for entry in VIPA_SOMMAIRE:
        if entry.section:
            story.append(Paragraph(pdf_safe(entry.section), st["h2"]))
        if entry.subsection:
            story.append(
                Paragraph(
                    pdf_safe(f"    {entry.subsection} ............... {entry.page}"),
                    st["body"],
                )
            )
    story.append(PageBreak())

    for sec in builder.all_sections():
        style = st["h1"] if sec.level == 1 else st["h2"]
        if sec.title:
            story.append(Paragraph(pdf_safe(sec.title), style))
        for para in sec.paragraphs:
            story.append(Paragraph(pdf_safe(para), st["body"]))
        for bullet in sec.bullets:
            story.append(Paragraph(pdf_safe(f"• {bullet}"), st["body"]))
        for tbl in sec.tables:
            _pdf_add_table(story, tbl, st["body"])
        if sec.level == 1:
            story.append(Spacer(1, 6))

    labels = [f"An {i + 1}" for i in range(HORIZON)]
    results = builder.results
    story.append(PageBreak())
    story.append(Paragraph(pdf_safe("Graphiques — comptes previsionnels"), st["h1"]))
    story.append(
        _pdf_bar_chart(
            "Chiffre d'affaires HT (TND)",
            labels,
            [year_val(results.revenue, y) for y in range(HORIZON)],
        )
    )
    story.append(Spacer(1, 10))
    story.append(
        _pdf_bar_chart(
            "Resultat net (TND)",
            labels,
            [year_val(results.netProfit, y) for y in range(HORIZON)],
            fill="#059669",
        )
    )
    story.append(Spacer(1, 10))
    story.append(
        _pdf_bar_chart(
            "Tresorerie cumulee (TND)",
            labels,
            [year_val(results.cumulativeTreasury, y) for y in range(HORIZON)],
            fill="#D97706",
        )
    )

    story.append(
        Paragraph(
            pdf_safe(
                f"Document généré par Business Plan Tunisie — "
                f"{datetime.now().strftime('%d/%m/%Y %H:%M')}"
            ),
            st["body"],
        )
    )
    doc.build(story)
    return str(out_path.resolve())


# —— XLSX ——


def _xlsx_header_row(ws, row: int = 1) -> None:
    fill = PatternFill("solid", fgColor="1E3A5F")
    font = Font(bold=True, color="FFFFFF")
    for cell in ws[row]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center")


def _xlsx_autowidth(ws, min_width: int = 10, max_width: int = 48) -> None:
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        width = min_width
        for cell in col:
            if cell.value is not None:
                width = max(width, min(len(str(cell.value)) + 2, max_width))
        ws.column_dimensions[letter].width = width


def _xlsx_write_section_sheet(ws, title: str, sections: list[SectionBlock]) -> None:
    ws.append([title])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])
    for sec in sections:
        if sec.title:
            ws.append([sec.title])
            ws[ws.max_row][0].font = Font(bold=True, size=11)
        for para in sec.paragraphs:
            ws.append([para])
        for bullet in sec.bullets:
            ws.append([f"• {bullet}"])
        for tbl in sec.tables:
            if tbl.caption:
                ws.append([tbl.caption])
            ws.append(tbl.headers)
            _xlsx_header_row(ws, ws.max_row)
            for row in tbl.rows:
                ws.append(row)
            ws.append([])
        ws.append([])


def build_feasibility_xlsx(
    plan_id: str,
    inputs: PlanInputs,
    results: PlanResults,
    path: Path,
    *,
    plan_title: str | None = None,
    extra_inputs: dict | None = None,
) -> str:
    builder = _make_builder(
        plan_id, inputs, results, plan_title=plan_title, extra_inputs=extra_inputs
    )
    wb = Workbook()

    # Sommaire
    ws0 = wb.active
    ws0.title = "Sommaire"
    ws0.append(["ETUDE DE FAISABILITE"])
    ws0["A1"].font = Font(bold=True, size=16)
    ws0.append([builder.company])
    ws0.append([plan_title or builder.company])
    ws0.append([f"Date : {datetime.now().strftime('%d/%m/%Y')}"])
    ws0.append([f"Référence : {plan_id[:8].upper()}"])
    ws0.append([])
    ws0.append(["SOMMAIRE"])
    ws0.append(["Section", "Rubrique", "Page"])
    _xlsx_header_row(ws0, 3)
    for entry in VIPA_SOMMAIRE:
        ws0.append([entry.section or "", entry.subsection or "", entry.page])
    _xlsx_autowidth(ws0)

    all_secs = builder.all_sections()
    part_labels = {
        "presentation": "Presentation projet",
        "marche": "Etude de marche",
        "invest": "Investissement",
        "finance": "Financement",
        "comptes": "Comptes previsionnels",
        "synthese": "Synthese conclusion",
    }
    for part, sheet_name in part_labels.items():
        secs = [s for s in all_secs if s.part == part]
        if not secs:
            continue
        ws = wb.create_sheet(sheet_name[:31])
        _xlsx_write_section_sheet(ws, sheet_name, secs)

    out = path / f"etude_faisabilite_{plan_id}_{_slug(builder.company)}.xlsx"
    wb.save(out)
    return str(out.resolve())
